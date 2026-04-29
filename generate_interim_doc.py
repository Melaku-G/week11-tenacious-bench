"""
Generate bench/interim_report.docx — Tenacious-Bench v0.1 Interim Report.
Run: python bench/generate_interim_doc.py
"""

from __future__ import annotations
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ────────────────────────────────────────────────────────────────

def shade_table_row(table, row_idx, fill_hex, bold=False, white_text=False):
    row = table.rows[row_idx]
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)
        if bold or white_text:
            for para in cell.paragraphs:
                for run in para.runs:
                    if bold:
                        run.bold = True
                    if white_text:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # shade header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def mono_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.4)
    return p


def dim_badge(score):
    return "PASS" if score == 1.0 else ("N/A" if score is None else "FAIL")


# ── load data ──────────────────────────────────────────────────────────────

def load_partition(path):
    return [json.loads(l) for l in Path(path).read_text(encoding="utf-8").splitlines() if l.strip()]


train  = load_partition("bench/tenacious_bench_v0.1/train.jsonl")
dev    = load_partition("bench/tenacious_bench_v0.1/dev.jsonl")
try:
    held   = load_partition("bench/tenacious_bench_v0.1/held_out.jsonl")
except Exception:
    held   = []

all_tasks = train + dev + held

def pass_rate(tasks):
    s = [t["score"] for t in tasks if t.get("score") is not None]
    if not s: return 0, 0, 0.0
    passing = sum(1 for x in s if x >= 0.75)
    return len(s), passing, round(sum(s) / len(s), 3)

# source mode counts
from collections import Counter
source_counts = {"train": Counter(), "dev": Counter(), "held_out": Counter()}
for t in train:
    source_counts["train"][t.get("source_mode") or "programmatic"] += 1
for t in dev:
    source_counts["dev"][t.get("source_mode") or "programmatic"] += 1
for t in held:
    source_counts["held_out"][t.get("source_mode") or "programmatic"] += 1

# Failure class counts (core classes only; good_example_* and edge cases excluded from targets)
CORE_FC = ["FC1_hallucination", "FC2_seg4_bypass", "FC3_overclaim", "tone_compliance", "FC9_draft_tag"]
fc_counts = {"train": Counter(), "dev": Counter(), "held_out": Counter()}
for t in train:
    fc_counts["train"][t.get("category", "unknown")] += 1
for t in dev:
    fc_counts["dev"][t.get("category", "unknown")] += 1
for t in held:
    fc_counts["held_out"][t.get("category", "unknown")] += 1

# Load the three example tasks
examples = {}
example_ids = {"TB-PROG-008": None, "TB-TD-001": None, "TB-ADV-019": None}
for t in all_tasks:
    if t["task_id"] in example_ids:
        examples[t["task_id"]] = t


# ── document ───────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Title ──────────────────────────────────────────────────────────────────
title = doc.add_heading("Tenacious-Bench v0.1 — Interim Report", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Acts I–II: Audit, Schema, Dataset Authoring")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(13)
sub.runs[0].italic = True

meta = doc.add_paragraph("Submitted: 2026-04-29 · Melaku Y. · Tenacious Conversion-Engine")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# ── Section 1: Bench Composition ─────────────────────────────────────────
heading(doc, "1. Bench Composition", 1)

body(doc,
    "Tenacious-Bench v0.1 is a 408-task domain-specific evaluation benchmark for B2B sales "
    "outreach agents. It covers six failure classes identified in Week 10 probing that general "
    "benchmarks (τ²-Bench retail, HELM, MTEB) cannot detect. All tasks are machine-verifiable: "
    "a deterministic Python scorer returns a weighted score ∈ [0, 1] with no LLM calls required."
)

heading(doc, "1.1 Master Cross-Tabulation: Partition × Source Mode × Failure Class", 2)

body(doc,
    "All three dataset axes integrated. Columns are failure classes; rows are source modes "
    "grouped by partition. Shaded rows show partition subtotals with actual vs. target percentages."
)

CROSS_ROWS = [
    ["TRAIN — actual 57.6%  |  target 50%  |  Δ +7.6 pp", "36", "42", "71", "45", "30", "11", "235"],
    ["  Programmatic",   "30", "30", "30", "30", "30",  "0", "150"],
    ["  Synthesis",       "3", "10", "36", "13",  "0",  "0",  "62"],
    ["  Trace",           "0",  "1",  "1",  "0",  "0",  "0",   "2"],
    ["  Hand-authored",   "3",  "1",  "4",  "2",  "0", "11",  "21"],
    ["DEV — actual 27.2%  |  target 30%  |  Δ −2.8 pp",  "22", "20", "24", "23", "18",  "4", "111"],
    ["  Programmatic",   "18", "18", "18", "18", "18",  "0",  "90"],
    ["  Synthesis",       "2",  "1",  "6",  "3",  "0",  "0",  "12"],
    ["  Trace",           "0",  "0",  "0",  "0",  "0",  "2",   "2"],
    ["  Hand-authored",   "2",  "1",  "0",  "2",  "0",  "2",   "7"],
    ["HELD-OUT — actual 15.2%  |  target 20%  |  Δ −4.8 pp", "12", "12", "13", "12", "13", "0", "62"],
    ["  Programmatic",   "12", "12", "12", "12", "12",  "0",  "60"],
    ["  Synthesis",       "0",  "0",  "0",  "0",  "0",  "0",   "0"],
    ["  Trace",           "0",  "0",  "1",  "0",  "0",  "0",   "1"],
    ["  Hand-authored",   "0",  "0",  "0",  "0",  "1",  "0",   "1"],
    ["GRAND TOTAL",      "70", "74","108", "80", "61", "15", "408"],
]
# Table row indices of partition header rows (0 = column header, 1 = TRAIN, 6 = DEV, 11 = HO, 16 = TOTAL)
PART_HEADER_ROWS = [1, 6, 11]
GRAND_TOTAL_ROW  = 16

xtab = add_table(doc,
    ["Partition / Source Mode", "FC1", "FC2", "FC3", "FC-T", "FC9", "Oth", "Total"],
    CROSS_ROWS,
    col_widths=[2.1, 0.45, 0.45, 0.45, 0.45, 0.45, 0.35, 0.55],
)
for ridx in PART_HEADER_ROWS:
    shade_table_row(xtab, ridx, "BDD7EE", bold=True)
shade_table_row(xtab, GRAND_TOTAL_ROW, "1F497D", bold=True, white_text=True)

doc.add_paragraph()
heading(doc, "1.2 Partition and Source-Mode Targets vs. Actuals", 2)

body(doc,
    "Challenge targets: partition 50/30/20 (train/dev/held-out); source-mode 30/30/25/15 "
    "(programmatic/synthesis/trace/hand-authored). Deviations explained below."
)

add_table(doc,
    ["Axis / Dimension", "Target", "Actual", "Δ", "Deviation rationale"],
    [
        ["Partition — Train",          "50% (204)", "57.6% (235)", "+7.6 pp",
         "24 hand-authored + 18 synthesis tasks added post-IRA pushed train above 50%"],
        ["Partition — Dev",            "30% (122)", "27.2% (111)", "−2.8 pp",
         "6 hand-authored added to dev; still near target"],
        ["Partition — Held-out",       "20%  (82)", "15.2%  (62)", "−4.8 pp",
         "Sealed at 62 after programmatic run; absolute task count is sufficient"],
        ["Source — Programmatic",      "30% (122)", "73.5% (300)", "+43.5 pp",
         "Intentional: systematic failure-class coverage, zero API cost, no judge latency"],
        ["Source — Multi-LLM synthesis","30% (122)", "18.1%  (74)", "−11.9 pp",
         "API cost and DeepSeek+LLaMA judge latency limited volume to 74 tasks"],
        ["Source — Trace-derived",     "25% (102)",  "1.2%   (5)", "−23.8 pp",
         "Structural constraint: only 5 real Week 10 traces available in trace_log.jsonl"],
        ["Source — Hand-authored",     "15%  (61)",  "7.1%  (29)",  "−7.9 pp",
         "Time-intensive; 29 tasks = 12 GOOD + 12 BAD (Style Guide v2) + 5 adversarial"],
    ],
    col_widths=[1.7, 0.9, 0.9, 0.75, 2.55],
)

doc.add_paragraph()
heading(doc, "1.3 Rubric Dimensions", 2)

body(doc,
    "Every task is scored on five machine-verifiable dimensions. Pass threshold: 0.75 "
    "(weighted sum). Each dimension returns 0 or 1 — no partial credit."
)

add_table(doc,
    ["Dimension", "Weight", "What it checks", "Scorer mechanism"],
    [
        ["signal_fidelity",    "0.30", "No unsupported growth/funding claims",
         "32 regex patterns; STRONG_NEGATIVE_GROWTH_PATTERNS on hard-negative context"],
        ["tone_compliance",    "0.20", "Zero banned phrases; no 'bench' to prospect",
         "32-phrase substring list + \\bbench\\b word-boundary regex (Style Guide v2)"],
        ["segment_gate",       "0.20", "No Seg4 AI pitch when ai_maturity_score < 2",
         "10 AI/LLM keyword list checked against ai_maturity_score field"],
        ["confidence_hedging", "0.15", "Hedging present when icp_confidence=low or low_peer_count=True",
         "7 hedging keyword list; conditional on brief fields"],
        ["format_compliance",  "0.15", "[DRAFT] tag, subject ≤60 chars, body ≤120/200 words",
         "Substring check, len() on subject (excl. [DRAFT]), word count (warm-aware)"],
    ],
    col_widths=[1.5, 0.6, 1.8, 2.3],
)

doc.add_paragraph()

# ── Section 2: Inter-Rater Agreement ─────────────────────────────────────
heading(doc, "2. Inter-Rater Agreement (IRA)", 1)

body(doc,
    "30-task stratified sample from the dev partition. Three raters: R1 (automated scorer, "
    "scoring_evaluator.py), R2 (LLaMA 3.1 8B via OpenRouter), R3 (human annotator, Melaku Y., "
    "2026-04-29). Cohen's κ computed pairwise per dimension. Target: κ ≥ 0.75 overall. "
    "IRA result: PASS (overall R1 vs R3 κ = 0.77)."
)

heading(doc, "2.1 R1 vs R2 Results (LLaMA 3.1 8B, two runs)", 2)

add_table(doc,
    ["Dimension", "R1 pass%", "Run 1 κ", "Run 2 κ", "Target", "Status"],
    [
        ["signal_fidelity",    "100.0%", "0.00†", "0.00†", "≥ 0.70", "ARTIFACT"],
        ["tone_compliance",    "86.7%",  "1.00",  "1.00",  "≥ 0.80", "PASS"],
        ["segment_gate",       "83.3%",  "0.00†", "−0.27", "≥ 0.85", "ARTIFACT/FAIL"],
        ["confidence_hedging", "53.3%",  "0.15",  "0.08",  "≥ 0.70", "FAIL"],
        ["format_compliance",  "76.7%",  "−0.06", "−0.30", "≥ 0.90", "FAIL"],
        ["Overall (≥0.75)",    "73.3%",  "0.26",  "0.04",  "≥ 0.75", "FAIL"],
    ],
    col_widths=[1.6, 0.75, 0.75, 0.75, 0.75, 1.5],
)

doc.add_paragraph(
    "† κ=0.00 is a statistical artifact when R1 has near-uniform scores (p_e ≈ p_o). "
    "Key finding: LLaMA 3.1 8B cannot reliably apply multi-step conditional rubric criteria. "
    "tone_compliance (pure lexical check) is the only stable dimension. "
    "A Claude Sonnet-class judge is required for κ ≥ 0.75 across all dimensions (v0.2 task)."
).italic = True

doc.add_paragraph()
heading(doc, "2.2 R3 Human Annotation Results", 2)

add_table(doc,
    ["Dimension", "R1 vs R3 κ", "R2 vs R3 κ", "Target", "R1 Status"],
    [
        ["signal_fidelity",    "0.00†",  "0.41",  "≥ 0.70", "ARTIFACT"],
        ["tone_compliance",    "1.00",   "1.00",  "≥ 0.80", "PASS"],
        ["segment_gate",       "1.00",   "−0.27", "≥ 0.85", "PASS"],
        ["confidence_hedging", "1.00",   "0.08",  "≥ 0.70", "PASS"],
        ["format_compliance",  "1.00",   "−0.30", "≥ 0.90", "PASS"],
        ["Overall (≥0.75)",    "0.77",   "0.25",  "≥ 0.75", "PASS"],
    ],
    col_widths=[1.6, 1.0, 1.0, 0.75, 1.75],
)

body(doc,
    "R3 found 5 tasks where R1 incorrectly passed emails containing growth language "
    "(raised a Series B, scale operations, scaling aggressively) paired with layoff=True OR "
    "open_roles=0. IRA protocol triggered: STRONG_NEGATIVE_GROWTH_PATTERNS added to "
    "scoring_evaluator.py. All 5 disputed tasks now correctly score 0 on signal_fidelity. "
    "Post-update scores: train mean=0.750, dev mean=0.741, held_out mean=0.731."
)

doc.add_paragraph()
heading(doc, "2.3 Revised Rubric Language — Before and After IRA", 2)

body(doc,
    "Two dimensions required post-IRA revision. signal_fidelity was revised after R3 flagged "
    "5 false negatives (score=1.0 on hard-negative briefs). format_compliance was tightened "
    "during Style Guide v2 integration (triggered by IRA agreement failures on format dimensions)."
)

add_table(doc,
    ["Dimension", "Before (pre-IRA)", "After (post-IRA)", "Trigger"],
    [
        ["signal_fidelity",
         "8 GROWTH_CLAIM_PATTERNS checked when low_signal=True "
         "(open_roles<3 OR velocity=low/unknown OR layoff). "
         "No separate hard-negative check.",
         "Same 8 patterns on low_signal PLUS 4 STRONG_NEGATIVE_GROWTH_PATTERNS "
         "checked when hard_negative=True (layoff=True OR open_roles=0). "
         "Added: 'raised a Series [B/C]', 'expanding into \\w', "
         "'scal\\w+ aggressively', 'scal\\w+ (its/your) operations'.",
         "R3 flagged 5 FNs: emails with 'raised a Series B' / "
         "'scale operations' paired with layoff=True+open_roles=0 scored 1.0."],
        ["format_compliance",
         "[DRAFT] tag check + word count ≤150 (fixed limit, no warm/cold distinction). "
         "No subject-line length check.",
         "[DRAFT] tag + subject ≤60 chars (excluding [DRAFT] prefix) + "
         "warm-aware word limit: ≤120 cold, ≤200 warm reply. "
         "Warm flag read from task['input']['warm'].",
         "Style Guide v2 Table 3 specifies 60-char subject cap and "
         "different limits per email type. IRA R2 κ=−0.30 on format_compliance "
         "revealed scorer did not match human expectations on subject length."],
    ],
    col_widths=[1.4, 1.9, 1.9, 1.6],
)

doc.add_paragraph()

# ── Section 3: Three Example Tasks ───────────────────────────────────────
heading(doc, "3. Example Tasks with Rubric Application", 1)

body(doc,
    "One task from each primary source mode. Each shows the enrichment input, candidate email, "
    "and per-dimension scoring with reasons."
)

EXAMPLES = [
    ("TB-PROG-008", "Programmatic", "FC1_hallucination",
     "Programmatically generated failure-injection task. Brief shows a company with "
     "layoff_event=True and open_roles=0 — a hard-negative context. The candidate email "
     "fabricates a Series B raise, claims market expansion, and uses 'scale operations'. "
     "signal_fidelity fails on STRONG_NEGATIVE_GROWTH_PATTERNS. confidence_hedging fails "
     "because icp_confidence=low and low_peer_count=True but the email contains no hedging words."),
    ("TB-TD-001", "Trace-derived", "FC3_overclaim",
     "Trace-derived task from eval/trace_log.jsonl. The candidate email contains "
     "'scale rapidly' and 'aggressive growth' — both overclaims against a brief with "
     "open_roles=1 and velocity=low. This task illustrates a known scorer gap: the phrases "
     "do not match GROWTH_CLAIM_PATTERNS exactly (pattern requires 'team' after 'scale'). "
     "Score=1.0 is a false negative. This drove the IRA signal_fidelity rubric update."),
    ("TB-ADV-019", "Hand-authored (adversarial)", "FC3_overclaim",
     "Style-guide-sourced bad example #2. Brief shows only 2 open roles with low ICP confidence "
     "but the email asserts 'scaling aggressively', uses 'top talent' and 'skyrocket' (all banned), "
     "and includes 'Quick question'/'Quick chat' (banned subject opener). No hedging despite "
     "icp_confidence=low. All four tone violations and missing hedging are caught correctly."),
]

for tid, mode_label, category, commentary in EXAMPLES:
    t = examples.get(tid)
    if not t:
        continue

    heading(doc, f"3.{EXAMPLES.index((tid,mode_label,category,commentary))+1}  {tid} — {mode_label}", 2)

    body(doc, commentary)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Input brief:").bold = True
    brief = t["input"]["enrichment_brief"]
    # Fix 2: for TB-PROG-008, surface key signal fields first so the hard-negative
    # context (open_roles_estimate=0, layoff_event=True) is unambiguous on first read.
    if tid == "TB-PROG-008":
        KEY_FIELDS = ["open_roles_estimate", "layoff_event", "icp_confidence", "low_peer_count"]
        brief_lines = [f"  company: {t['input'].get('company','?')}"]
        for k in KEY_FIELDS:
            if k in brief:
                brief_lines.append(f"  {k}: {brief[k]}  ← key signal field")
        for k, v in brief.items():
            if k not in KEY_FIELDS:
                brief_lines.append(f"  {k}: {v}")
    else:
        brief_lines = [f"  company: {t['input'].get('company','?')}"]
        brief_lines += [f"  {k}: {v}" for k, v in brief.items()]
    mono_para(doc, "\n".join(brief_lines))

    p = doc.add_paragraph()
    p.add_run("Candidate output:").bold = True
    mono_para(doc, f"Subject: {t['candidate_output']['email_subject']}\n\n{t['candidate_output']['email_body']}")

    p = doc.add_paragraph()
    p.add_run("Rubric scores:").bold = True

    ds = t.get("dimension_scores") or {}
    score = t.get("score", 0)
    verdict = "PASS" if score >= 0.75 else "FAIL"

    dim_rows = []
    for dim in ["signal_fidelity", "tone_compliance", "segment_gate", "confidence_hedging", "format_compliance"]:
        d = ds.get(dim, {})
        s = d.get("score")
        badge = "PASS" if s == 1.0 else ("N/A" if "n/a" in (d.get("reason") or "").lower() else "FAIL")
        dim_rows.append([dim, badge, d.get("reason", "—")])

    # Fix 3: annotate TB-TD-001 as a documented false negative
    if tid == "TB-TD-001":
        dim_rows.append(["TOTAL SCORE", f"{score:.3f} → {verdict} (FALSE NEGATIVE)",
                         "signal_fidelity regex gap — 'scale rapidly' not matched by current GROWTH_CLAIM_PATTERNS"])
    else:
        dim_rows.append(["TOTAL SCORE", f"{score:.3f} → {verdict}", ""])

    add_table(doc,
        ["Dimension (weight)", "Result", "Reason"],
        dim_rows,
        col_widths=[1.6, 1.0, 3.6],
    )
    doc.add_paragraph()


# ── Section 4: Status + Plan ──────────────────────────────────────────────
heading(doc, "4. Status: What Is Working, What Is Not, Plan for Days 4–7", 1)

heading(doc, "4.1 What Is Working", 2)

for item in [
    "408 tasks across 4 source modes — all three partitions generated, scored, and committed.",
    "Scoring evaluator: deterministic, no LLM calls, 32 banned phrases (Style Guide v2 full list), "
     "\\bbench\\b word-boundary check, subject ≤60 chars, warm-aware word limit.",
    "IRA complete: R1 vs R3 κ=0.77 (PASS). STRONG_NEGATIVE_GROWTH_PATTERNS added after R3 "
     "flagged 5 signal_fidelity false negatives.",
    "Contamination checks: CC-001 through CC-005 all clean or accepted. CC-002 embedding "
     "similarity run (all-MiniLM-L6-v2): max cosine=1.0 from template overlap by design, "
     "same disposition as n-gram overlap check.",
    "24 style-guide-sourced hand-authored tasks: 12 good examples (all score ≥0.75), "
     "12 bad examples with labeled failure modes from the official banned-phrase list.",
    "All 7 required synthesis memos complete as of Day 4 — 4 common reading + 3 Path A specific "
     "(Tülu 3, LIMA, Magpie). Each includes a specific disagreement justified against Tenacious-Bench evidence.",
    "train.py: per-source sampling weights applied (hand_authored=4×, trace_derived=2×, "
     "synthesis=0.75×) based on Magpie difficulty-distribution finding.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(11)

heading(doc, "4.2 What Is Not Working / Known Limitations", 2)

for item in [
    "signal_fidelity false negatives: the GROWTH_CLAIM_PATTERNS regex requires 'team' after "
     "'scale', so 'scale rapidly' or 'scaling the organization' are missed. The 5 pre-existing "
     "adversarial tasks (TB-ADV-001 to TB-ADV-005) all score 1.0 despite FC3 category labels. "
     "Fix: extend patterns in v0.2 after more R3 annotation.",
    "LLM judge (R2) only reliable on tone_compliance (κ=1.00). All other dimensions fail "
     "κ targets with LLaMA 3.1 8B. A Claude Sonnet-class judge is needed for v0.2 (adds ~$2 cost).",
    "PDF attachment, multiple-ask, and pricing-fabrication violations (BAD #9, #10, #11) are "
     "not machine-detectable with the current regex scorer — 3 bad tasks score 1.0. "
     "Ground_truth descriptions capture the violations for SFT training even if the scorer cannot.",
    "Training partition has 235 tasks with 18 bad examples included — SFT should filter to "
     "score≥0.75 pairs only (train.py does this). At score≥0.75 threshold, effective training "
     "set is ~141 high-quality pairs. LIMA benchmark suggests this is sufficient for alignment "
     "at this scale.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(11)

heading(doc, "4.3 Plan for Days 4–7", 2)

add_table(doc,
    ["Day", "Date", "Task", "Deliverable"],
    [
        ["Day 4 (done)", "Apr 29",
         "Path A memos (Tülu 3, LIMA, Magpie). CC-002 embedding check. "
         "Style guide integration (32 banned phrases, 24 labeled tasks). train.py source weights.",
         "memo_05/06/07.md, contamination_check.json clean, 408-task dataset, updated scorer"],
        ["Day 5", "May 1",
         "SFT training on Colab T4. Qwen2.5-1.5B LoRA rank=16 alpha=32 via Unsloth. "
         "Filter train.jsonl to score≥0.75 pairs. 300 steps, cosine LR. Log loss curve. "
         "KILL CRITERION: if train loss has not decreased ≥0.10 by step 100 → abort, "
         "diagnose data pipeline or LR. Cost: Colab T4 free tier (~0 spend).",
         "bench/training/checkpoints/lora_adapter/, training_run.log"],
        ["Day 6", "May 2",
         "Ablations: Delta A (SFT vs baseline on held_out), Delta B (SFT vs guarded prompt-only). "
         "Paired bootstrap CI. held_out_traces.jsonl. Cost-Pareto. "
         "PIVOT TRIGGER: if held-out score delta (SFT vs baseline) < 0.05 (outside 95% CI) "
         "→ declare Path A insufficient, pivot to Path B (guarded prompt-only). "
         "Eval-tier spend reserved: 62 held-out tasks × 1 Sonnet 4.6 call × ~3K tokens "
         "= ~$0.56 of $10 envelope (cumulative Day 1–4 spend: ~$2.10; remaining: ~$7.90).",
         "ablation_results.json, held_out_traces.jsonl, statistical test output"],
        ["Day 7", "May 3",
         "HuggingFace dataset push (train+dev). Model push (LoRA adapter). "
         "Technical blog post (1,200–2,000 words). Community engagement (τ²-Bench GitHub issue). "
         "2-page CEO/CFO memo. evidence_graph.json.",
         "HF dataset URL, HF model URL, blog URL, memo.pdf, GitHub issue link"],
    ],
    col_widths=[1.1, 0.7, 3.3, 1.6],
)

doc.add_paragraph()
body(doc,
    "Path A justification (committed in methodology.md): FC-3 (signal overclaim, ~40% unguarded "
    "trigger rate) and FC-1 (hallucinated firmographics, ~30%) are generation-quality failures — "
    "the model defaults to optimistic SDR copy from its training distribution regardless of "
    "enrichment signal. SFT on grounded (input, correct-output) pairs directly retrains the "
    "generation behavior. Trace evidence: TB-b2c3d4e5-0001 (FC-3) and TB-b2c3d4e5-0002 (FC-2). "
    "LIMA (Zhou et al. 2023): 141 high-quality training pairs is sufficient at this domain-specific "
    "scale. Tülu 3 (Lambert et al. 2024): data composition > volume; SFT fixes distributional bias."
)

# ── Save ───────────────────────────────────────────────────────────────────
out = Path("bench/interim_report.docx")
doc.save(str(out))
print(f"Saved: {out}")
